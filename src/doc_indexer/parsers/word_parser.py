"""
Word document parser with LLM support.
"""

import io
from pathlib import Path
from typing import List

from docx import Document as DocxDocument
from PIL import Image

from doc_indexer.parsers.base import BaseParser, PageContent


class WordParser(BaseParser):
    """Parser for Word documents with image extraction support."""

    def __init__(self, *args, extract_images: bool = True, **kwargs) -> None:
        """
        Initialize Word parser.

        Args:
            extract_images: Whether to extract embedded images
            *args, **kwargs: Arguments passed to BaseParser
        """
        super().__init__(*args, **kwargs)
        self.extract_images = extract_images

    async def extract_pages(self, file_path: Path) -> List[PageContent]:
        """Extract content from Word document, treating sections as pages."""
        doc = DocxDocument(str(file_path))
        pages = []

        current_page_text = []
        current_page_tables = []
        current_page_images = []
        page_number = 1

        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                current_page_text.append(paragraph.text)

            if self._has_page_break(paragraph):
                pages.append(
                    self._create_page_content(
                        page_number,
                        current_page_text,
                        current_page_tables,
                        current_page_images,
                    )
                )

                current_page_text = []
                current_page_tables = []
                current_page_images = []
                page_number += 1

        for table in doc.tables:
            table_data = self._extract_table_data(table)
            current_page_tables.append(table_data)

        if self.extract_images:
            images = self._extract_images_from_docx(doc)
            current_page_images.extend(images)

        if current_page_text or current_page_tables or current_page_images:
            pages.append(
                self._create_page_content(
                    page_number,
                    current_page_text,
                    current_page_tables,
                    current_page_images,
                )
            )

        if not pages:
            all_text = []
            all_tables = []

            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    all_text.append(paragraph.text)

            for table in doc.tables:
                all_tables.append(self._extract_table_data(table))

            pages.append(
                PageContent(
                    page_number=1,
                    text="\n".join(all_text),
                    image=current_page_images[0] if current_page_images else None,
                    tables=all_tables,
                )
            )

        return pages

    def _has_page_break(self, paragraph) -> bool:
        """
        Check if a paragraph contains a page break.

        Args:
            paragraph: Docx paragraph object

        Returns:
            True if page break found
        """
        # Check for explicit page breaks in the paragraph
        for run in paragraph.runs:
            if "w:br" in run._element.xml and 'type="page"' in run._element.xml:
                return True

        # Check for section breaks
        if paragraph._element.getnext() is not None:
            next_elem = paragraph._element.getnext()
            if next_elem.tag.endswith("sectPr"):
                return True

        return False

    def _create_page_content(
        self,
        page_number: int,
        text_parts: List[str],
        tables: List[dict],
        images: List[Image.Image],
    ) -> PageContent:
        """
        Create a PageContent object from collected content.

        Args:
            page_number: Page number
            text_parts: List of text paragraphs
            tables: List of table data
            images: List of extracted images

        Returns:
            PageContent object
        """
        return PageContent(
            page_number=page_number,
            text="\n".join(text_parts) if text_parts else "",
            image=images[0] if images else None,  # Use first image for page
            tables=tables,
        )

    def _extract_table_data(self, table) -> dict:
        """
        Extract data from a Word table.

        Args:
            table: Docx table object

        Returns:
            Dictionary with table data
        """
        table_data = {"header": [], "rows": []}

        for i, row in enumerate(table.rows):
            row_data = []
            for cell in row.cells:
                row_data.append(cell.text.strip())

            if i == 0:
                # Assume first row is header
                table_data["header"] = row_data
            else:
                table_data["rows"].append(row_data)

        return table_data

    def _extract_images_from_docx(self, doc: DocxDocument) -> List[Image.Image]:
        """
        Extract embedded images from Word document.

        Args:
            doc: Docx document object

        Returns:
            List of PIL Images
        """
        images = []

        try:
            # Access the document's relationships
            for rel in doc.part.rels.values():
                if "image" in rel.reltype:
                    try:
                        # Get image data
                        image_data = rel.target_part.blob

                        # Convert to PIL Image
                        image = Image.open(io.BytesIO(image_data))

                        # Resize if too large
                        if image.width > 1920 or image.height > 1080:
                            image.thumbnail((1920, 1080), Image.Resampling.LANCZOS)

                        images.append(image)
                    except Exception as e:
                        print(f"Error extracting image: {e}")
        except Exception as e:
            print(f"Error accessing document images: {e}")

        return images
